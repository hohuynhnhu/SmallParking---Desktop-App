import pandas as pd
from docx import Document
from docx.shared import Inches

# -------------------------
# 1. Đọc dữ liệu từ CSV
# -------------------------
df = pd.read_csv(r"D:\license-plate-recognition\codequetbienso\training_metrics.csv")

# -------------------------
# 2. Tạo document Word
# -------------------------
doc = Document()
doc.add_heading("Bảng Training Metrics CNN", level=1)

# -------------------------
# 3. Thêm bảng vào Word
# -------------------------
# Tạo bảng với số dòng +1 (dòng header) và số cột đúng với dataframe
table = doc.add_table(rows=df.shape[0]+1, cols=df.shape[1])
table.style = 'Table Grid'  # kiểu bảng có đường viền rõ

# Thêm header
for j, col_name in enumerate(df.columns):
    table.cell(0, j).text = str(col_name)

# Thêm dữ liệu
for i in range(df.shape[0]):
    for j in range(df.shape[1]):
        value = df.iloc[i, j]
        # format float với 3 chữ số thập phân
        if isinstance(value, float):
            value = f"{value:.3f}"
        table.cell(i+1, j).text = str(value)

# -------------------------
# 4. Lưu file Word
# -------------------------
output_path = r"D:\license-plate-recognition\codequetbienso\training_metrics.docx"
doc.save(output_path)
print(f"✅ File Word đã được lưu: {output_path}")
