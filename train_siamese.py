import os
import random
import torch
import torch.nn as nn
import torchvision.transforms as transforms
from torch.utils.data import Dataset, DataLoader
from PIL import Image
from torchvision.models import resnet18, ResNet18_Weights
from ultralytics import YOLO
import cv2
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

# === Hyperparameters ===
BATCH_SIZE = 8
LR = 1e-4
EPOCHS = 30
IMG_SIZE = (128, 128)
DATASET_ROOT = 'Dataset_Siamese'   # gồm train/, val/, test/
DEVICE = torch.device("cuda" if torch.cuda.is_available() else "cpu")

# === Transform ảnh ===
transform = transforms.Compose([
    transforms.Resize(IMG_SIZE),
    transforms.ToTensor(),
    transforms.Normalize(mean=[0.485, 0.456, 0.406],
                         std=[0.229, 0.224, 0.225])
])

# === Load YOLO (dùng để crop xe) ===
yolo_model = YOLO("yolov8n.pt")

# === Dataset Siamese ===
class SiameseDataset(Dataset):
    def __init__(self, dataset_dir, transform=None, num_pairs=2000):
        self.dataset_dir = dataset_dir
        self.transform = transform
        self.folders = [os.path.join(dataset_dir, f) for f in os.listdir(dataset_dir) if os.path.isdir(os.path.join(dataset_dir, f))]
        self.num_pairs = num_pairs

    def __len__(self):
        return self.num_pairs

    def crop_car_largest(self, img_path):
        results = yolo_model(img_path, imgsz=640, verbose=False)
        max_area = 0
        best_crop = None
        for r in results:
            for box in r.boxes:
                cls = int(box.cls[0])
                if yolo_model.names[cls] == "car":  # chỉ giữ car
                    x1, y1, x2, y2 = map(int, box.xyxy[0])
                    area = (x2 - x1) * (y2 - y1)
                    if area > max_area:
                        img = cv2.imread(img_path)
                        crop = img[y1:y2, x1:x2]
                        if crop is not None and crop.size > 0:
                            best_crop = Image.fromarray(cv2.cvtColor(crop, cv2.COLOR_BGR2RGB))
                            max_area = area
        if best_crop is not None:
            return best_crop
        else:
            return Image.open(img_path).convert("RGB")

    def __getitem__(self, idx):
        same_class = random.choice([True, False])  # 50% cùng xe

        if same_class:
            folder = random.choice(self.folders)
            images = os.listdir(folder)
            if len(images) < 2:
                return self.__getitem__(idx)
            img1_name, img2_name = random.sample(images, 2)
            label = 0
            img1_path = os.path.join(folder, img1_name)
            img2_path = os.path.join(folder, img2_name)
        else:
            folder1, folder2 = random.sample(self.folders, 2)
            img1_list = os.listdir(folder1)
            img2_list = os.listdir(folder2)
            if not img1_list or not img2_list:
                return self.__getitem__(idx)
            img1_path = os.path.join(folder1, random.choice(img1_list))
            img2_path = os.path.join(folder2, random.choice(img2_list))
            label = 1

        img1 = self.crop_car_largest(img1_path)
        img2 = self.crop_car_largest(img2_path)

        if self.transform:
            img1 = self.transform(img1)
            img2 = self.transform(img2)

        return img1, img2, torch.tensor(label, dtype=torch.float32)

# === Mạng trích đặc trưng ===
class SiameseNetwork(nn.Module):
    def __init__(self):
        super(SiameseNetwork, self).__init__()
        self.backbone = resnet18(weights=ResNet18_Weights.IMAGENET1K_V1)
        self.backbone.fc = nn.Linear(self.backbone.fc.in_features, 128)

    def forward_once(self, x):
        return self.backbone(x)

    def forward(self, x1, x2):
        return self.forward_once(x1), self.forward_once(x2)

# === Loss Contrastive ===
class ContrastiveLoss(nn.Module):
    def __init__(self, margin=1.0):
        super().__init__()
        self.margin = margin

    def forward(self, out1, out2, label):
        euclidean_distance = nn.functional.pairwise_distance(out1, out2)
        loss = torch.mean((1 - label) * torch.pow(euclidean_distance, 2) +
                          label * torch.pow(torch.clamp(self.margin - euclidean_distance, min=0.0), 2))
        return loss

# === Hàm tính loss cho 1 dataloader ===
def evaluate(model, dataloader, criterion):
    model.eval()
    total_loss = 0
    with torch.no_grad():
        for img1, img2, label in dataloader:
            img1, img2, label = img1.to(DEVICE), img2.to(DEVICE), label.to(DEVICE)
            out1, out2 = model(img1, img2)
            loss = criterion(out1, out2, label)
            total_loss += loss.item()
    return total_loss / len(dataloader)

# === Huấn luyện + Xuất báo cáo ===
def train():
    train_dataset = SiameseDataset(os.path.join(DATASET_ROOT, "train"), transform=transform, num_pairs=8000)
    val_dataset   = SiameseDataset(os.path.join(DATASET_ROOT, "val"), transform=transform, num_pairs=2000)
    test_dataset  = SiameseDataset(os.path.join(DATASET_ROOT, "test"), transform=transform, num_pairs=2000)

    train_loader = DataLoader(train_dataset, shuffle=True, batch_size=BATCH_SIZE)
    val_loader   = DataLoader(val_dataset, shuffle=True, batch_size=BATCH_SIZE)
    test_loader  = DataLoader(test_dataset, shuffle=True, batch_size=BATCH_SIZE)

    model = SiameseNetwork().to(DEVICE)
    criterion = ContrastiveLoss()
    optimizer = torch.optim.Adam(model.parameters(), lr=LR)

    train_history = []
    val_history = []

    for epoch in range(EPOCHS):
        model.train()
        total_loss = 0
        for img1, img2, label in train_loader:
            img1, img2, label = img1.to(DEVICE), img2.to(DEVICE), label.to(DEVICE)
            out1, out2 = model(img1, img2)
            loss = criterion(out1, out2, label)
            optimizer.zero_grad()
            loss.backward()
            optimizer.step()
            total_loss += loss.item()

        avg_train_loss = total_loss / len(train_loader)
        avg_val_loss = evaluate(model, val_loader, criterion)

        train_history.append(avg_train_loss)
        val_history.append(avg_val_loss)

        print(f"Epoch {epoch+1}/{EPOCHS}, Train Loss: {avg_train_loss:.4f}, Val Loss: {avg_val_loss:.4f}")

    # Lưu model
    torch.save(model.state_dict(), "siamese_model.pth")
    print("✅ Model đã được lưu thành siamese_model.pth")

    # Test sau khi train
    test_loss = evaluate(model, test_loader, criterion)
    print(f"📊 Test Loss: {test_loss:.4f}")

    # === Vẽ biểu đồ Train vs Val Loss ===
    plt.figure(figsize=(6,4))
    plt.plot(range(1, EPOCHS+1), train_history, marker="o", label="Train Loss")
    plt.plot(range(1, EPOCHS+1), val_history, marker="s", label="Val Loss")
    plt.xlabel("Epoch")
    plt.ylabel("Loss")
    plt.title("Training & Validation Loss per Epoch")
    plt.legend()
    plt.grid(True)
    plt.savefig("loss_chart.png")
    plt.close()

    # === Xuất báo cáo DOCX ===
    doc = Document()
    doc.add_heading("Báo cáo huấn luyện Siamese Network", 0)

    doc.add_heading("1. Thông tin huấn luyện", level=1)
    doc.add_paragraph(f"Số epoch: {EPOCHS}")
    doc.add_paragraph(f"Batch size: {BATCH_SIZE}")
    doc.add_paragraph(f"Learning rate: {LR}")
    doc.add_paragraph(f"Số cặp ảnh train: {len(train_dataset)}")
    doc.add_paragraph(f"Số cặp ảnh val: {len(val_dataset)}")
    doc.add_paragraph(f"Số cặp ảnh test: {len(test_dataset)}")

    doc.add_heading("2. Biểu đồ Loss", level=1)
    doc.add_picture("loss_chart.png", width=Inches(5))

    doc.add_heading("3. Bảng thống kê Loss", level=1)
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Epoch'
    hdr_cells[1].text = 'Train Loss'
    hdr_cells[2].text = 'Val Loss'
    for i, (tr, vl) in enumerate(zip(train_history, val_history), 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = f"{tr:.4f}"
        row_cells[2].text = f"{vl:.4f}"

    doc.add_heading("4. Kết quả trên Test set", level=1)
    doc.add_paragraph(f"Test Loss: {test_loss:.4f}")

    doc.save("Training_Report.docx")
    print("📄 Báo cáo đã được lưu thành Training_Report.docx")

if __name__ == '__main__':
    train()
