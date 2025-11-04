

import argparse
import os
import glob
import time
from pathlib import Path
import csv
import math

import pandas as pd
import matplotlib.pyplot as plt
from ultralytics import YOLO
from docx import Document
from docx.shared import Inches

# --------------------
# Helper utils
# --------------------
def safe_float(x):
    """Try convert to float, fallback to string."""
    try:
        return float(x)
    except Exception:
        try:
            return float(x.item())
        except Exception:
            return str(x)

def find_latest_run(project_dir):
    """Find newest subfolder in runs/train or project directory."""
    p = Path(project_dir)
    if not p.exists():
        return None
    subs = [d for d in p.iterdir() if d.is_dir()]
    if not subs:
        return None
    subs_sorted = sorted(subs, key=lambda d: d.stat().st_mtime, reverse=True)
    return subs_sorted[0]

def read_results_csv(run_dir):
    """Try to read ultralytics results.csv (or epoch_metrics.csv fallback)."""
    run_dir = Path(run_dir)
    candidates = [
        run_dir / "results.csv",
        run_dir / "metrics.csv",
        run_dir / "epoch_metrics.csv",  # our callback fallback
        run_dir / "logs_epoch_metrics.csv",
    ]
    for c in candidates:
        if c.exists():
            try:
                return pd.read_csv(c), c
            except Exception:
                pass
    # try glob for results*.csv
    for c in run_dir.glob("*.csv"):
        try:
            df = pd.read_csv(c)
            # heuristics: must have 'epoch' column or numeric columns
            if "epoch" in df.columns or df.shape[1] > 1:
                return df, c
        except Exception:
            continue
    return None, None

# --------------------
# Callback to log per-epoch
# --------------------
def make_epoch_logger():
    """
    Return a callback function to append epoch metrics to CSV inside trainer.save_dir.
    The trainer argument is passed by ultralytics to callback.
    """
    def on_train_epoch_end(trainer):
        # trainer object gives access to many attributes: epoch, tloss, metrics, lr, save_dir, epoch_time ...
        epoch_idx = getattr(trainer, "epoch", None)
        if epoch_idx is None:
            return
        epoch = int(epoch_idx) + 1

        # get labeled train losses if available
        try:
            train_losses = trainer.label_loss_items(trainer.tloss, prefix="train") or {}
        except Exception:
            train_losses = {}

        # validation metrics (trainer.metrics is commonly a dict like {'val/mAP50': x, ...})
        try:
            val_metrics = dict(getattr(trainer, "metrics", {}) or {})
        except Exception:
            val_metrics = {}

        # learning rate
        try:
            lr = dict(getattr(trainer, "lr", {}) or {})
        except Exception:
            lr = {}

        epoch_time = getattr(trainer, "epoch_time", None)

        # convert all values to floats/strings
        row = {"epoch": epoch, "epoch_time": safe_float(epoch_time)}
        # train losses
        for k, v in train_losses.items():
            # remove 'train/' prefix if present
            key = k.replace("train/", "") if isinstance(k, str) else str(k)
            row[f"train_{key}"] = safe_float(v)
        # val metrics (keys may contain '/')
        for k, v in val_metrics.items():
            key = str(k).replace("/", "_")
            row[f"val_{key}"] = safe_float(v)
        # lr
        for k, v in lr.items():
            row[f"lr_{k}"] = safe_float(v)

        # determine CSV path
        save_dir = getattr(trainer, "save_dir", None)
        if save_dir is None:
            save_dir = Path("runs/train/exp")
        else:
            save_dir = Path(save_dir)

        save_dir.mkdir(parents=True, exist_ok=True)
        csv_path = save_dir / "epoch_metrics.csv"

        # write header if not exist
        write_header = not csv_path.exists()
        with open(csv_path, "a", newline="", encoding="utf-8") as f:
            writer = csv.DictWriter(f, fieldnames=list(row.keys()))
            if write_header:
                writer.writeheader()
            writer.writerow(row)

    return on_train_epoch_end

# --------------------
# Plotting utilities
# --------------------
def plot_metrics(df, save_dir):
    """
    df: pandas DataFrame with epoch rows and numeric columns
    For each numeric column except 'epoch', plot vs epoch and save PNG.
    """
    save_dir = Path(save_dir)
    save_dir.mkdir(parents=True, exist_ok=True)
    # ensure epoch column numeric
    if "epoch" in df.columns:
        df = df.sort_values("epoch")
        epochs = df["epoch"].values
    else:
        epochs = list(range(1, len(df) + 1))

    # select numeric columns except epoch
    numeric_cols = []
    for c in df.columns:
        if c == "epoch":
            continue
        if pd.api.types.is_numeric_dtype(df[c]):
            numeric_cols.append(c)

    plots = []
    for col in numeric_cols:
        plt.figure()
        plt.plot(epochs, df[col].values)
        plt.xlabel("Epoch")
        plt.ylabel(col)
        plt.title(col)
        plt.grid(True)
        out = save_dir / f"{col.replace(' ', '_').replace('/', '_')}.png"
        plt.savefig(out, bbox_inches="tight")
        plt.close()
        plots.append(out)
    return plots

# --------------------
# Generate docx report
# --------------------
def save_docx_report(csv_df, per_class_df, plot_paths, save_dir, hyperparams):
    doc = Document()
    doc.add_heading("YOLOv8 Training Report", level=1)
    doc.add_paragraph(f"Run directory: {save_dir}")
    doc.add_paragraph(f"Generated: {time.strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_heading("Hyperparameters", level=2)
    p = doc.add_paragraph()
    for k, v in hyperparams.items():
        p.add_run(f"{k}: ").bold = True
        p.add_run(f"{v}\n")
    doc.add_heading("Per-epoch metrics (sample)", level=2)

    # put first 200 rows or all if small
    rows_to_put = csv_df if len(csv_df) <= 200 else csv_df.head(200)
    table = doc.add_table(rows=1, cols=len(rows_to_put.columns))
    hdr_cells = table.rows[0].cells
    for i, c in enumerate(rows_to_put.columns):
        hdr_cells[i].text = str(c)
    for _, r in rows_to_put.iterrows():
        cells = table.add_row().cells
        for i, c in enumerate(rows_to_put.columns):
            cells[i].text = str(r[c])

    if per_class_df is not None:
        doc.add_page_break()
        doc.add_heading("Per-class AP", level=2)
        # small table
        table = doc.add_table(rows=1, cols=len(per_class_df.columns))
        hdr_cells = table.rows[0].cells
        for i, c in enumerate(per_class_df.columns):
            hdr_cells[i].text = str(c)
        for _, r in per_class_df.iterrows():
            cells = table.add_row().cells
            for i, c in enumerate(per_class_df.columns):
                cells[i].text = str(r[c])

    # insert plots
    doc.add_page_break()
    doc.add_heading("Plots", level=2)
    for p in plot_paths:
        doc.add_paragraph(p.name)
        try:
            doc.add_picture(str(p), width=Inches(6))
        except Exception as e:
            doc.add_paragraph(f"Could not add image {p.name}: {e}")

    out_doc = Path(save_dir) / "training_report.docx"
    doc.save(out_doc)
    return out_doc

# --------------------
# Main flow
# --------------------
def main(args):
    # Create model
    model = YOLO(args.model)

    # Register callback to log epoch metrics
    epoch_logger = make_epoch_logger()
    model.add_callback("on_train_epoch_end", epoch_logger)

    # Train
    train_kwargs = dict(
        data=args.data,
        epochs=args.epochs,
        imgsz=args.imgsz,
        batch=args.batch,
        project=args.project,
        name=args.name,
        exist_ok=True,
        val=True,  # ensure validation occurs during training
        device=args.device
    )
    print("▶ Starting training with args:", train_kwargs)
    results = model.train(**train_kwargs)  # this runs training and triggers callbacks

    # Locate run directory (prefer project/name)
    if args.project and args.name:
        run_dir = Path(args.project) / args.name
    else:
        # fallback to latest folder in runs/train
        run_dir = find_latest_run("runs/train")
    if run_dir is None:
        run_dir = Path("runs") / "train" / "exp"
    print("Run directory:", run_dir.resolve())

    # Read results CSV (ultralytics saves results.csv)
    df_results, csv_path = read_results_csv(run_dir)
    if df_results is None:
        print("⚠️ Không tìm thấy results.csv trong run dir, sẽ dùng epoch_metrics.csv nếu có.")
    else:
        print("Loaded results CSV:", csv_path)

    # If we have results from ultralytics, convert to pandas DataFrame (if not already)
    if df_results is not None and not isinstance(df_results, pd.DataFrame):
        try:
            df_results = pd.DataFrame(df_results)
        except Exception:
            pass

    # If no results.csv but epoch CSV exists
    epoch_csv = run_dir / "epoch_metrics.csv"
    if df_results is None and epoch_csv.exists():
        df_results = pd.read_csv(epoch_csv)

    if df_results is None:
        print("❌ Không có dữ liệu metrics (results.csv / epoch_metrics.csv). Bỏ qua phần vẽ biểu đồ.")
        df_results = pd.DataFrame()

    # Plot metrics
    if not df_results.empty:
        plots = plot_metrics(df_results, run_dir)
        print("Saved plots:", plots)
    else:
        plots = []

    # Run full validation to get per-class AP and confusion matrix
    try:
        print("▶ Running final validation to get per-class metrics...")
        val_res = model.val(data=args.data, plots=False)  # returns a Results-like object or dict
        # try to extract per-class AP table
        per_class_df = None
        # Many versions: val_res.summary() returns list/dict; val_res.curves or val_res.maps also possible
        try:
            # Try summary -> list of dicts
            summary = val_res.summary()  # may return list of dicts or string
            # If summary is list-like of dicts, convert
            if isinstance(summary, (list, tuple)):
                per_class_df = pd.DataFrame(summary)
        except Exception:
            per_class_df = None

        # fallback: if val_res has 'metrics' dict or 'results_dict'
        if per_class_df is None:
            try:
                rd = getattr(val_res, "results_dict", None)
                if rd:
                    per_class_df = pd.DataFrame([rd])
            except Exception:
                per_class_df = None

        # final fallback: try results.confusion_matrix summary
        if per_class_df is None:
            try:
                cm = getattr(val_res, "confusion_matrix", None)
                if cm is not None:
                    cm_summary = cm.summary(normalize=True, decimals=5)
                    per_class_df = pd.DataFrame(cm_summary)
            except Exception:
                per_class_df = None

    except Exception as e:
        print("⚠️ model.val() failed or returned unexpected structure:", e)
        per_class_df = None

    # Save docx report
    hyperparams = {
        "model": args.model,
        "data": args.data,
        "epochs": args.epochs,
        "imgsz": args.imgsz,
        "batch": args.batch,
        "device": args.device,
    }
    out_docx = save_docx_report(df_results, per_class_df, plots, run_dir, hyperparams)
    print("Saved report:", out_docx.resolve())


if __name__ == "__main__":
    
    parser = argparse.ArgumentParser()
    parser.add_argument("--data", type=str, required=True, help="path to data.yaml")
    parser.add_argument("--model", type=str, default="yolov8n.pt", help="base model or path")
    parser.add_argument("--epochs", type=int, default=50)
    parser.add_argument("--imgsz", type=int, default=640)
    parser.add_argument("--batch", type=int, default=16)
    parser.add_argument("--project", type=str, default="runs/train")
    parser.add_argument("--name", type=str, default="logo_experiment")
    parser.add_argument("--device", type=str, default="0", help="CUDA device or 'cpu'")
    args = parser.parse_args()
    main(args)
